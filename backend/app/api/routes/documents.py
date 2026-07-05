# 담당: 최수인 — 법령/개념 문서 저장 (문서 DB 구축, PRD 3.2 P0)
import io
import zipfile

from docx import Document as DocxFile
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from pydantic import BaseModel
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.db.models import Document
from app.db.session import get_db

router = APIRouter(prefix="/api/documents", tags=["documents"])

SUPPORTED_EXTENSIONS = {"pdf", "docx", "txt"}


def _extension(filename: str) -> str:
    return filename.lower().rsplit(".", 1)[-1] if "." in filename else ""


def _extract_text(filename: str, raw: bytes) -> str:
    ext = _extension(filename)
    if ext == "pdf":
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if ext == "docx":
        doc = DocxFile(io.BytesIO(raw))
        return "\n".join(p.text for p in doc.paragraphs).strip()
    if ext == "txt":
        return raw.decode("utf-8", errors="ignore").strip()
    raise ValueError(f"지원하지 않는 파일 형식입니다: .{ext}")


class DocumentCreate(BaseModel):
    subject: str
    title: str
    content: str
    source: str | None = None


class DocumentOut(BaseModel):
    id: int
    subject: str
    title: str
    content: str
    source: str | None

    class Config:
        from_attributes = True


@router.post("", response_model=DocumentOut)
def create_document(payload: DocumentCreate, db: Session = Depends(get_db)) -> Document:
    document = Document(**payload.model_dump())
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


@router.post("/upload", response_model=list[DocumentOut])
def upload_document(
    subject: str = Form(...),
    title: str | None = Form(None),
    source: str | None = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
) -> list[Document]:
    filename = file.filename or "문서"
    ext = _extension(filename)
    raw = file.file.read()

    # (파일명, 원본 바이트) 목록: 단일 파일이면 1개, ZIP이면 내부 파일 여러 개
    entries: list[tuple[str, bytes]] = []

    if ext == "zip":
        try:
            zf = zipfile.ZipFile(io.BytesIO(raw))
        except zipfile.BadZipFile:
            raise HTTPException(status_code=400, detail="ZIP 파일을 읽을 수 없습니다.")
        for name in zf.namelist():
            if name.endswith("/") or name.startswith("__MACOSX"):
                continue
            if _extension(name) not in SUPPORTED_EXTENSIONS:
                continue
            entries.append((name, zf.read(name)))
        if not entries:
            raise HTTPException(
                status_code=400,
                detail="ZIP 안에 PDF/DOCX/TXT 파일이 없습니다.",
            )
    elif ext in SUPPORTED_EXTENSIONS:
        entries.append((filename, raw))
    else:
        raise HTTPException(
            status_code=400,
            detail="PDF, DOCX, TXT, ZIP 파일만 업로드할 수 있습니다.",
        )

    documents: list[Document] = []
    for name, content_bytes in entries:
        try:
            text = _extract_text(name, content_bytes)
        except Exception:
            continue  # ZIP 내 개별 파일이 손상/미지원이면 건너뛰고 나머지는 계속 처리
        if not text:
            continue
        base_name = name.rsplit("/", 1)[-1].rsplit(".", 1)[0]
        doc_title = title if (title and len(entries) == 1) else base_name
        documents.append(Document(subject=subject, title=doc_title, content=text, source=source))

    if not documents:
        raise HTTPException(
            status_code=400,
            detail="텍스트를 추출하지 못했습니다. 스캔된 이미지이거나 빈 파일일 수 있습니다.",
        )

    db.add_all(documents)
    db.commit()
    for document in documents:
        db.refresh(document)
    return documents


@router.get("", response_model=list[DocumentOut])
def list_documents(subject: str | None = None, db: Session = Depends(get_db)) -> list[Document]:
    query = db.query(Document)
    if subject:
        query = query.filter(Document.subject == subject)
    return query.order_by(Document.id.desc()).limit(100).all()
